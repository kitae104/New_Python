# -*- coding: utf-8 -*-

from PyQt5.QtWidgets import QFileDialog
from selenium import webdriver
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from openpyxl import load_workbook
from email import encoders
from bs4 import BeautifulSoup
import smtplib
import urllib.request
import openpyxl
import xlrd
import operator
import docx2txt
from docx import Document
from PyQt5.QtWidgets import *
import time
import urllib.request 
from imagesearch import *
import pyautogui
import sys
import re
from docx.shared import Cm
import csv
from urllib.request import urlopen
import cv2
from PIL import Image
import pandas as pd
from collections import Counter
import matplotlib.pyplot as plt
import pytesseract
import konlpy.tag
from wordcloud import WordCloud


class Activity:

    driver = webdriver
    id = ''
    class_name = ''
    sid = ''
    part = ''
    attachfile = 0
    wb = ''
    ws = ''
    Name = ''
    list = ''
    document = ''
    rdr = ''
    wr = ''
    f = ''

    def __init__(self):
        self.__activity_dict = {  # 기능들의 이름, 변수, 변수값을 기본값으로 저장
            '브라우저 열기': {
                'url': '_url'
            },
            'id값으로 클릭': {
                'id': '_id'
            },
            'id값으로 입력': {
                'id': '_id', 'key': '_key'
            },
            'class값으로 클릭': {
                'class_name': '_class_name'
            },
            'class값으로 입력': {
                'Class': '_Class', 'key': '_key'
            },
            'css값으로 클릭': {
                'Css': '_Css'
            },
            'css값으로 입력': {
                'Css': '_Css', 'key': '_key'
            },
            'xpath값으로 클릭': {
                'XPath': '_XPath'
            },
            'xpath값으로 입력': {
                'XPath': '_XPath', 'key': '_key'
            },
            'Gmail 세션 생성': {
                'id': '_id', 'smtppw': '_smtppw'
            },
            '파일첨부': {
                'filename': '_filename'
            },
            '엑셀 읽어 보내기': {
                'filename': '_filename'
            },
            'text 크롤링': {
                'class_name': 'class_name'
            },
            'url 크롤링': {
                'class_name': 'class_name', 'company_class': '_company_class'
            },
            '텍스트 파일 읽기': {
                'filename': '_filename'
            },
            '텍스트 파일 쓰기': {
                'text': '_text'
            },
            '텍스트 추가': {
                'filename': '_filename', 'text': '_text'
            },
            '텍스트 파일을 엑셀 파일로 변환': {
                'text': '_text'
            },
            '시트 불러오기': {
                'filename': '_filename'
            },
            '행 삭제': {
                'row': '_row'
            },
            '열 삭제': {
                'col': '_col'
            },
            '셀 옮기기': {
                'cell': '_cell', 'move': '_move'
            },
            '값 얻기': {
                'cell': '_cell'
            },
            '행 찾기': {
                'row': '_row', 'text': '_text', 'totalrow': '_totalrow'
            },
            '값 설정': {
                'cell': '_cell', 'text': '_text'
            },
            'docx 읽기': {
                'filename': '_filename'
            },
            'docx 생성': {
                'style': '_style'
            },
            '이미지 추가': {
                'filename': '_filename', 'size': '_size'
            },
            'docx 작성': {
                'style': '_style'
            },
            'head 쓰기': {
                'text': '_text', 'style': '_style'
            },
            'body 쓰기': {
                'text': '_text', 'style': '_style'
            },
            'docx 저장': {

            },
            'csv 읽기': {
                'filename': '_filename'
            },
            '열 가져오기': {
                'selline': '_selline'
            },
            'csv 쓰기': {
                'text': '_text'
            },
            '행 추가': {
                'text': '_text'
            },
            'csv파일 추가모드로 열기': {
                'filename': '_filename'
            },
            '파일 닫기': {

            },
            '모든 값 얻기': {

            },
            '브라우저 닫기': {

            },
            '팝업창 닫기': {

            },
            '이미지 다운로드(src)': {
                'src': '_src'
            },
            '페이지 스크린샷': {

            },
            '이미지 클릭': {
                'ImageFile': '_imagefilename'
            },
            'url 여러 페이지 크롤링': {
                'class_name': '_class_name', 'company_class': '_company_class', 'count': '_count'
            },
            '데이터로 크롤링': {

            },
            '좌표 클릭': {
                'X': 'Pos_X', 'Y': 'Pos_Y'
            },
            '이미지크롤링': {
                'fordername':'_fordername' , 'filetype' : '_filetype' 
            },
            '딜레이': {
                'time': '_time'
            },
            '반복_For': {
                'startnum': '_startnum', 'endnum': '_endnum', 'roofnum': '_roofnum'
            },

            '이미지 사이즈 변환':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path','X':'_X','Y':'_Y'
            },

            '이미지 사이즈 정사각형 변환':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path'
            },
            '이미지 회색조 변환':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path'
            },
            '이미지 이진화 변환':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path'
            },
            '이미지 히스토그램 균일화(칼라)':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path'
            },
            '이미지 히스토그램 균일화(회색조)':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path'
            },
            '이미지 가우시안 블러':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path','Kernerl_X':'_Kernerl_X','Kernerl_Y':'_Kernerl_Y','Sigma' : '_Sigma'
            },
            '이미지 미디안 블러':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path','Kernerl_Size':'_Kernerl_Size'
            },
            '이미지 양방향 블러':{
                'Original_Path':'_Original_Path','Save_Path':'_Save_Path','Kernerl_Size':'_Kernerl_Size','Sigma_Space' : '_Sigma_Space','Sigma_Color' : '_Sigma_Color'
            },
            '텍스트 결측치 행 전체 삭제':{
                'filename':'_filename'
            },
            '텍스트 결측치 열 전체 삭제':{
                'filename':'_filename'
            },
            '텍스트 결측치 사용자값 삽입':{
                'filename':'_filename','text':'_text'
            },
            '텍스트 결측치 액셀 확인':{
                'filename':'_filename'
            },
            '텍스트 결측치 직전값으로 삽입':{
                'filename':'_filename'
            },
            '텍스트 결측치 직후값으로 삽입':{
                'filename':'_filename'
            },
            '텍스트 결측치 평균값으로 삽입':{
                'filename':'_filename'
            },
            '텍스트 결측치 중앙값으로 삽입':{
                'filename':'_filename'
            },
            '텍스트 결측치 최빈값으로 삽입':{
                'filename':'_filename'
            },
            'Kor_Ocr':{
                'fordername':'_fordername'
            },
            'Eng_Ocr':{
                'fordername':'_fordername'
            },
            '막대 그래프 시각화':{
                'fordername':'_fordername'
            },
            '텍스트 불용어 처리':{
                'filename':'_filename'
            },
            '텍스트 형태소 분석':{
                'filename':'_filename'
            },
            '텍스트 명사 추출':{
                'filename':'_filename'
            },
            '워드 클라우드 시각화':{
                'filename':'_filename'
            },
            '꺾은선 그래프 시각화':{
                'filename':'_filename'
            },
        }

    # key의 갯수 세기
    def count_keys(self, name):
        keylist = list(self.__activity_dict[name].keys())  # key값들 저장
        return len(keylist)

    # 번호에 해당하는 key의 값
    def search_name_key(self, name, num):
        keylist = list(self.__activity_dict[name].keys())   # key값들 저장
        return keylist[num]

    # key에 해당하는 value
    def search_key_value(self, name, key):
        value = self.__activity_dict[name][key]  # key에 해당하는 value
        return value

    def OpenBrowser(self, url):
        driver = webdriver.Chrome()  # 크롬 웹 드라이버 생성
        driver.maximize_window()  # 페이지창 최대화
        driver.implicitly_wait(3)  # 3초대기
        driver.get(url)  # url주소로 이동
        self.driver = driver  # 드라이버 전역변수로 저장

    def find_element_by_id_click(self, id):
        self.id = self.driver.find_element_by_id(id)
        self.id.click()

    def find_element_by_id_sendkey(self, id, key):
        self.driver.find_element_by_id(id).send_keys(key)

    def find_element_by_class_click(self, _class):
        self.class_name = self.driver.find_element_by_class_name(_class)
        self.class_name.click()

    def find_element_by_class_sendkey(self, _class, key):
        self.driver.find_element_by_class_name(_class).send_keys(key)

    def find_element_by_css_click(self, Css):
        self.driver.find_element_by_css_selector(Css).click()

    def find_element_by_css_sendkey(self, Css, key):
        self.driver.find_element_by_css_selector(Css).send_keys(key)

    def find_element_by_xpath_click(self, XPath):
        self.driver.find_element_by_xpath(XPath).click()

    def find_element_by_xpath_sendkey(self, XPath, key):
        self.driver.find_element_by_xpath(XPath).send_keys(key)

    def Gmail_Session_Create(self, id, pw):  # STMP세션 생성(GMAIL)
        self.session = smtplib.SMTP("smtp.gmail.com", 587)  # 세션생성
        self.session.starttls()  # SMTP 연결을 TLS (전송 계층 보안) 모드로 전환, 그뒤에 모든 STMP명령은 암호화되기 위함
        self.sid = id  # STMP id저장
        self.session.login(id, pw)  # SMTP 계정 인증

    def Send_Mail(self, name, receiver, subject, content):  # 메일 전송

        self.msg = MIMEMultipart()  # MIME 메세지 생성

        self.msg["From"] = self.sid  # 메세지의 송신자 설정
        self.msg["To"] = receiver  # 메세지의 수신자 저장
        self.msg["Subject"] = subject  # 메세지의 내용 저장

        text = MIMEText(_text=content)
        self.msg.attach(text)  # 내용 첨부
        if self.attachfile == 1:  # 파일을 첨부할 경우 : attachfile == 1
            self.msg.attach(self.part)  # part에 저장된 파일 첨부
        self.session.sendmail(self.sid, receiver, self.msg.as_string())  # 메일 전송(STMPid, 메세지의 수신자, 메세지의 내용)

    def Attach_File(self, Filename):  # 파일 첨부
        path = Filename  # 파일이름 저장
        with open(path, 'rb') as f:  # 파일 열기
            self.part = MIMEBase("application", "octet-stream")  # 이진 파일 형식으로 변환
            self.part.set_payload(f.read())  # 7계층 프로토콜 적용
            encoders.encode_base64(self.part)  # 영상, 이미지 파일을 문자열 형태로 변환
            self.part.add_header('Content-Disposition', 'attachment', filename=path)  # Content-Disposition헤더를 attachment속성으로 명시하여 파일 첨부를 활성화
            self.attachfile = True  # 파일이 있음을 의미

    def Excel_read_send(self, excelFile):  # 엑셀 파일 읽은 후 전송 함수 호출
        wb = load_workbook(excelFile)  # 엑셀파일 불러와 워크북 생성
        ws = wb.active  # 워크시트 생성
        for row in ws.iter_rows():  # 열단위로 읽어 name, receiver, subject, content에 저장
            name = row[1].value
            receiver = row[2].value
            subject = row[3].value
            content = row[4].value
            Activity.Send_Mail(self, name, receiver, subject, content)  # 전송 함수 호출
        self.session.quit()  # 세션 종료

    def WebCrawling_text(self, _class):  # 크롤링해서 text를 읽고 txt파일로 저장
        # 파일 지정해서 열기
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        f = open(file_name, "w", encoding="utf-8")

        # 크롤링
        h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함 (서버 측에서 자동 프로그램이 동작하지 않게 하기위해 막아 놓은거 푸는 기능)
        url = self.driver.current_url  # 브라우저의 현재 url을 가져온다.
        req = urllib.request.Request(url, headers=h)  # 벤 걸리지 않게 함
        reqhttp = urllib.request.urlopen(req)
        http = reqhttp.read()
        soup = BeautifulSoup(http, 'html.parser')

        class_choice = soup.find_all(class_=_class)  # 크롤링할 class에서 가져온 값을 class_choice에 저장

        for count in class_choice:
            text = str(count.get_text().strip()) + "\n"  # 텍스트를 단위별로 읽어와서 text에 태그들 제거하고 저장한다.
        text = text.replace("// flash 오류를 우회하기 위한 함수 추가", "")
        text = text.replace("function _flash_removeCallback() {}", "")
        f.write(text)  # text에 저장된 내용을 file에 쓴다.
        f.close()

    def WebCrawling_url(self, _class, _company_class):  # 크롤링해서 text를 읽고 text에 저장된 하이퍼링크를 가져와서 엑셀 파일로 저장
        wb = openpyxl.Workbook()  # 워크북 생성
        ws = wb.active  # 워크시트 생성

        # 크롤링
        h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함
        url = self.driver.current_url
        print(url)
        req = urllib.request.Request(url, headers=h)  # 벤 걸리지 않게 함
        reqhttp = urllib.request.urlopen(req)
        http = reqhttp.read()  # .decode('utf-8') #읽고 한글 깨지지 않게 함
        soup = BeautifulSoup(http, 'html.parser')

        class_choice = soup.find_all(class_=_class)  # 크롤링할 class에서 가져온 값을 class_choice에 저장
        data = []
        text = []
        company = []
        for n in class_choice:
            textsave = str(n.get_text().strip())
            datasave = n["href"]  # 크롤링해서 가져온 정보를 href로 변환해서 datasave에 저장
            data.append(datasave)  # datasave에 저장된 값을 순서대로 data리스트에 저장
            text.append(textsave)  # textsave에 저장된 값을 순서대로 text리스트에 저장

        company_class_choice = soup.find_all(class_=_company_class)
        for n in company_class_choice:
            datasave = str(n.get_text().strip())
            datasave = datasave.replace("언론사 선정", "")
            company.append(datasave)

        for data_range in range(len(data)):  # data에 저장된 개수만큼 반복문을 돌림
            ws.cell(row=data_range + 1, column=1).value = data[data_range]  # 엑셀 파일에 순서대로 저장 (1,1), (2,1) ....
            ws.cell(row=data_range + 1, column=2).value = text[data_range]
            ws.cell(row=data_range + 1, column=3).value = company[data_range]

        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        wb.save(name[0] + '.xlsx')

    def WebCrawling_url_count(self, _class, _company_class, _count):  # 페이지 수를 지정하여 크롤링함.(WebCrawling_url과 페이지 수 지정을 제외하면 동일)
        wb = openpyxl.Workbook()  # 워크북 생성
        ws = wb.active  # 워크시트 생성

        # 크롤링
        h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함
        first_url = self.driver.current_url  # 현재 url을 가져옴
        req = urllib.request.Request(first_url, headers=h)  # 벤 걸리지 않게 함
        reqhttp = urllib.request.urlopen(req)
        http = reqhttp.read()  # .decode('utf-8') #읽고 한글 깨지지 않게 함
        soup = BeautifulSoup(http, 'html.parser')

        # 페이지 수 지정
        pageNum = 1
        lastPage = int(_count) * 10 - 8
        data = []  # 엑셀에 저장되어 있는 url을 저장할 변수 생성
        text = []  # 엑셀에 저장되어 있는 제목을 저장할 변수 생성
        company = []  # 엑셀에 저장되어 있는 회사 이름을 저장할 변수 생성

        while pageNum < lastPage:  # 지정된 페이지 수 보다 작을 경우 반복문 실행
            pageCheck = "&start=" + str(pageNum) + "&refresh_start=0"  # 페이지 수를 지정할 내용
            url = first_url + pageCheck  # 가져온 url뒤에 페이지 수에 대한 정보를 이어 붙임
            html = urllib.request.urlopen(url).read()
            soup = BeautifulSoup(html, 'html.parser')

            class_choice = soup.find_all(class_=_class)  # 신문기사 제목의 클래스 이름(제목만 크롤링하면, 제목과 url을 가져올 수 있음)
            for n in class_choice:
                textsave = str(n.get_text().strip())
                datasave = n["href"]  # 크롤링해서 가져온 정보를 href로 변환해서 datasave에 저장
                data.append(datasave)  # datasave에 저장된 값을 순서대로 data에 저장
                text.append(textsave)  # textsave에 저장된 값을 순서대로  text에 저장

            company_class_choice = soup.find_all(class_=_company_class) # 신문기사 회사의 클래스 이름
            for n in company_class_choice:
                datasave = str(n.get_text().strip()) # 회사의 이름을 읽음
                datasave = datasave.replace("언론사 선정", "")  # 네이버에서 언론사 선정이라는 값을 추가로 입력할 때가 있기 때문에, 회사이름만 크롤링하기 위해 "언론사 선정" 제거
                company.append(datasave)  # datasave에 저장된 값을 company에 저장

            pageNum += 10  # 네이버에서 다음 페이지로 넘어갈 때, 필요한 값

        for data_range in range(len(data)):  # data, text, company에 저장된 값을 엑셀에 저장
            ws.cell(row=data_range + 1, column=1).value = data[data_range]
            ws.cell(row=data_range + 1, column=2).value = text[data_range]
            ws.cell(row=data_range + 1, column=3).value = company[data_range]

        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        wb.save(name[0] + '.xlsx')

    def WebCrawling_get_url(self):  # WebCrawling_url, WebCrawling_url_count를 사용하여 생성된 엑셀파일을 읽어와서 그 값으로 크롤링하는 함수
        name = QFileDialog.getOpenFileName()  # 파일 열기
        wb = xlrd.open_workbook(name[0])  # 엑셀파일 열기
        ws = wb.sheet_by_index(0)  # 첫번째 시트
        numrow = ws.nrows  # 열의 수 numrow에 저장

        data_url = []
        for row in range(numrow):
            if (ws.row_values(row)[0]):  # row에 값이 있다면(url)
                text = ws.row_values(row)[0]  # 그 값을 text에 저장
                data_url.append(text)  # text에 저장된 값을 data_url에 저장

        data_text = []
        for row in range(numrow):
            if (ws.row_values(row)[1]):  # row에 값이 있다면(제목)
                text = ws.row_values(row)[1]  # 그 값을 text에 저장
                data_text.append(text)  # text에 저장된 값을 data_text에 저장

        data_company = []
        for row in range(numrow):
            if (ws.row_values(row)[2]):  # row에 값이 있다면(회사 이름)
                text = ws.row_values(row)[2]  # 그 값을 text에 저장
                data_company.append(text)  # text에 저장된 값을 data_company에 저장

        dirname = QFileDialog.getExistingDirectory()  # 크롤링한 txt파일을 저장할 폴더 지정

        for num in range(len(data_url)):
            file_name = data_text[num] + ".txt"  # 신문기사 제목으로 txt파일 이름 지정
            file_name = re.sub('[-=+,#/\?:^$@*\"※~&%ㆍ!』"\\‘|\(\)\[\]\<\>`\'》]', '', file_name)  # 제목에 특수문자가 존재한다면 제거
            file_name = dirname + '/' + file_name  # txt파일 저장할 폴더와 이름 지정
            f = open(file_name, "w", encoding="utf-8")
            # 크롤링
            h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함 (서버 측에서 자동 프로그램이 동작하지 않게 하기위해 막아 놓은거 푸는 기능)
            url = data_url[num]
            req = urllib.request.Request(url, headers=h)  # 벤 걸리지 않게 함
            reqhttp = urllib.request.urlopen(req)
            http = reqhttp.read()
            soup = BeautifulSoup(http, 'html.parser')

            # 회사별로 크롤링
            if data_company[num] == '연합뉴스':
                class_choice = soup.find_all(class_='story-news article')
            elif data_company[num] == '파이낸셜뉴스':
                class_choice = soup.find_all(class_='article_cont mg cont_txt_read')
            elif data_company[num] == '뉴시스':
                class_choice = soup.find_all(id_='textBody')
            elif data_company[num] == '스포츠서울':
                class_choice = soup.find_all(class_='news_text')
            elif data_company[num] == '뉴스핌':
                class_choice = soup.find_all(id_='news_contents')
            elif data_company[num] == '헬스조선':
                class_choice = soup.find_all(class_='par')
            elif data_company[num] == '머니투데이':
                class_choice = soup.find_all(class_='view_content')
            elif data_company[num] == '한국경제TV':
                class_choice = soup.find_all(class_='box-news-body')
            elif data_company[num] == '서울경제':
                class_choice = soup.find_all(class_='article_view')
            elif data_company[num] == 'SBS':
                class_choice = soup.find_all(class_='text_area')
            elif data_company[num] == '조세일보':
                class_choice = soup.find_all(articleBody_='articleBody')
            elif data_company[num] == '오마이뉴스':
                class_choice = soup.find_all(class_='at_contents')
            elif data_company[num] == '뉴스1':
                class_choice = soup.find_all(class_='detail sa_area')
            elif data_company[num] == '노컷뉴스':
                class_choice = soup.find_all(class_='articleBody')
            elif data_company[num] == '한국경제':
                class_choice = soup.find_all(id_='articletxt')
            elif data_company[num] == 'KBS':
                class_choice = soup.find_all(class_='detail-body font-size')
            elif data_company[num] == '스포츠동아':
                class_choice = soup.find_all(class_='article_word')
            elif data_company[num] == 'ZDNet Korea':
                class_choice = soup.find_all(class_='view_cont')
            elif data_company[num] == 'IT조선':
                class_choice = soup.find_all(class_='article font03')
            elif data_company[num] == '동아일보':
                class_choice = soup.find_all(class_='article_txt')
            elif data_company[num] == '국민일보':
                class_choice = soup.find_all(class_='tx')
            elif data_company[num] == '블로터':
                class_choice = soup.find_all(class_='article--content')
            elif data_company[num] == 'MBC':
                class_choice = soup.find_all(class_='news_txt')
            elif data_company[num] == '헤럴드경제':
                class_choice = soup.find_all(class_='con_left_bd')
            elif data_company[num] == '경향신문':
                class_choice = soup.find_all(class_='art_body')
            else:
                class_choice = soup.find_all(class_='')

            for class_count in class_choice:
                text = str(class_count.get_text().strip()) + "\n"  # 텍스트를 단위별로 읽어와서 text에 태그들 제거하고 저장한다.
            text = text.replace("// flash 오류를 우회하기 위한 함수 추가", "")
            text = text.replace("function _flash_removeCallback() {}", "")
            f.write(text)  # text에 저장된 내용을 file에 쓴다.
            f.close()

    def Read_TextFile(self, fName):  # 파일이름으로 txt 파일 열기
        f = open(fName, "r", encoding="utf8")  # 텍스트 파일 read 형태로 열기
        self.list = f.readlines()  # list에 txt파일 모든 줄 저장하기

    def Write_TextFile(self, text):  # 새로운 txt파일쓰기
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"  # 다른이름으로 파일 저장
        f = open(file_name, "w+t", encoding="utf-8")  # 위에서 저장한 파일 쓰기 형태로 열기
        f.write(text+'\n')  # text에 입력된 값을 txt에 쓴다.

    def Add_Text(self, filename, text):  # 기존 txt에 추가하기
        f = open(filename, "at", encoding='utf-8')  # txt파일을 수정 형태로 연다.
        f.writelines(text+'\n')  # 입력된 값을 한줄씩 입력한다.

    def txt_to_xlsx(self, Division):  # txt 내에서 데이터를 구별할 값 입력
        wb = openpyxl.Workbook()  # sheet 만들기
        sheet = wb.active  # sheet 활성화 하기
        for i in range(1, len(self.list) + 1):  # 1~리스트의 크기 + 1의 범위 까지 반복
            list1 = self.list[i - 1].split(Division)  # 리스트 0부터 division으로 입력된 값으로 글자 자르기
            for k in range(1, len(list1) + 1):  # 1~리스트1의 크기 +1 만큼 반복
                sheet[chr(k + 64) + str(i)].value = list1[k - 1]  # ASCII의 65부터 A,B,...시트 번호 A1부터 하나씩 입력한다.
        name = QFileDialog.getSaveFileName()
        wb.save(name[0] + '.xlsx')
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Load_Sheet(self, fName):  # 이용할 엑셀파일을 워크북으로 로딩하기
        self.wb = openpyxl.load_workbook(fName)  # sheet를 load
        self.ws = self.wb.active
        self.Name = fName  # 파일네임 그대로 저장

    def Delete_rows(self, row):  # 열 지우기
        rowNum = int(row)  # 입력받은 문자를 int로 바꿔주기
        self.ws.delete_rows(rowNum)  # row 지우기
        self.wb.save(self.Name)  # 파일 저장
        QMessageBox.about(self, "message", "삭제되었습니다.")

    def Delete_cols(self, col):  # 행 지우기
        colNum = int(col)  # 입력받은 문자를 int로 바꿔주기
        self.ws.delete_cols(colNum)  # col 지우기
        self.wb.save(self.Name)  # 파일저장
        QMessageBox.about(self, "message", "삭제되었습니다.")

    def Move_Cell(self, Range, move_range):  # 이동시킬 값의 범위, 얼만큼 이동할란지? ex)3,3
        word_list = move_range.split(',')  # 입력받은 범위를 ,로 split해 list에 각 문자 담기
        moverow = int(word_list[0])  # 리스트의 문자열을 숫자로 바꿔주기
        movecol = int(word_list[1])
        self.ws.move_range(Range, rows=moverow, cols=movecol, translate=True)
        self.wb.save(self.Name)
        QMessageBox.about(self, "message", "완료되었습니다.")

    def GetValue(self, cell):  # 지정범위 값 txt 파일로 적기
        div = list(cell)  # 전달받은 범위를 한글자씩 쪼개서 리스트에 담기
        cell_range = self.ws[cell]  # 입력받은 셀범위를 워크시트에서 찾기
        a = list()
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        f = open(file_name, "w", encoding="utf-8")
        if ":" in div:  # 여러셀 가져오기
            for row in cell_range:  # 셀의 개수만큼 반복
                for cell in row:  # row만큼 반복
                    a.append(cell.value)  # 리스트에 값 추가
            for i in range(len(a)):
                f.write('%s\n' % a[i])  # txt파일에 적기
        else:  # 지정한 셀만 필요시
            f.write(cell_range.value)

    def Find_row(self, col , text, totalcol):  # 지정한col에 text가 존재하는 열의 값들 모두 가져오기
        colnum = int(col)
        a = list()
        n = int(totalcol)
        for i in self.ws.rows:  # 모든 row를 검사
            if operator.eq(i[colnum - 1].value, text):  # 해당 컬럼에서 text와 같은 값이 존재하나?
                for j in range(0, n):  # 총 col개수만큼 반복
                    a.append(i[j].value)  # 리스트에 저장
        result = [a[i * n:(i + 1) * n] for i in range((len(a) + n - 1) // n)]  # 리스트에 셀의 값 하나씩 저장
        num = int(len(a) / n)   # a리스트의 크기를 총 col의 개수로 나눈 값을 num에 저장
        self.wb.close()
        wb2 = openpyxl.Workbook()  # 값 저장 할 wb 열기
        ws2 = wb2.active
        for i in range(1, num + 1):
            for j in range(0, n):
                ws2.cell(row=i, column=j + 1).value = result[i - 1][j]  # ws2의 셀에 값 입력
        name = QFileDialog.getSaveFileName()
        wb2.save(name[0] + '.xlsx')

    def SetValue(self, cell, text):  # 지정한 셀에 값 넣기
        self.ws[cell] = text
        self.wb.save(self.Name)
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Read_docx(self, fileName):  # 워드파일 내용 읽어오기
        text = docx2txt.process(fileName)  # 워드파일 읽어오기
        QMessageBox.about(self, "message", text)

    def Make_docx(self,style):
        if 'docx' in style:  # style변수에 스타일이 지정된 워드파일이 입력됐을 경우
            self.document = Document(style)  # document를 지정한 스타일로 만들기
        else:
            self.document = Document()  # document를 지정한 형식 없이 만들기

    def Add_Image(self,filename,size):
        if '_size' in size :  # size에 입력값을 안넣었을 경우
            self.document.add_picture(filename)
        else:
            size2 = size.split(',')  # 사이즈를 ,로 구분
            sizelist2 = [int(x) for x in size2]  # size를 리스트에 저장
            self.document.add_picture(filename,width=Cm(sizelist2[0]),height= Cm(sizelist2[1]))  # cm단위로 사진 붙이기

    def Write_head(self, text, style):  # 제목작성
        if '_style' in style:  # style에 아무것도 지정되지 않았을 경우
            self.document.add_heading(text, 0)  # 헤더의 크기를 기본 0 으로 지정하거나 , Make_docx에서 style을 지정했을 경우엔 위에서 지정한 값으로 지정
        else:
            numstyle = int(style)  # 입력받은 값을 int로 바꿔주기
            self.document.add_heading(text, numstyle)  # 기본 워드의 스타일 번호로 스타일을 준다.

    def Write_body(self, text, style):  # 본문작성
        if style == '_style':
            self.document.add_paragraph(text)  # style에 지정된 값이 없을 경우엔 기본으로 값 입력
        else:
            self.document.add_paragraph(text, style=style)  # style이 지정된 경우 스타일을 준다.

    def Save_docx(self):  # 워드파일 저장
        name = QFileDialog.getSaveFileName()
        self.document.save(name[0] + '.docx')

    def Read_csv(self,filename):  # 읽기전용으로 csv파일 열기
        f = open(filename, 'r', encoding='euc-kr')  # csv파일 읽기
        self.rdr = csv.reader(f)  # csv파일 리더 생성

    def Get_Line(self, selectline):
        linenum = int(selectline)
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        txt = open(file_name, "w", encoding='utf-8')  # 텍스트파일 쓰기형식으로 열기
        for line in self.rdr:
            txt.write(line[linenum] + '\n')  # 텍스트파일에 해당줄의 값 입력
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Open_Add_mode(self,filename):
        self.f = open(filename, 'a', newline='')  # 수정형식으로 파일 열기
        self.wr = csv.writer(self.f)  # writer지정

    def Add_row(self,text):
        rs2 = text.split(',')  # 입력된 값을 ,로 구분
        self.wr.writerow(rs2)  # row입력
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Write_csv(self, text):
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".csv"
        f = open(file_name, "w", encoding='euc-kr', newline='')  # 쓰기전용으로 파일열기
        wr = csv.writer(f)  # writer 지정
        rs2 = text.split(',')
        wr.writerow(rs2)

    def File_Close(self):
        self.f.close()

    def Get_All_Value(self):  # 모든 값 가져와 txt 파일로 저장
        a = list()
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        text = open(file_name, "w", encoding='utf-8')  # 쓰기형식으로 파일열기
        for txt in self.rdr:
            a.append(txt)  # 리스트에 텍스트 추가
        for i in range(len(a)):
            text.write('%s\n' % a[i])  # txt파일에 값입력
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Close_Browser(self): # 브라우저 닫는 함수
        self.driver.close() # 현재 사용중이던 브라우저를 닫는다.

    def Close_PopUp(self): # 브라우저에 생성된 팝업 창을 닫는 함수
        driver_count = len(self.driver.window_handles) # 현재 열려 있는 창의 수를 driver_count에 저장
        if driver_count != 1: # 열려있는 창이 1개가 아니라면(팝업 창이 존재한다면)
            driver_count = driver_count - 1
            self.driver.switch_to.window(self.driver.window_handles[driver_count])  # 팝업창을 선택
            self.driver.close()
        self.driver.switch_to.window(self.driver.window_handles[0])  # main창 선택

    def Image_Download(self, src):
        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        urllib.request.urlretrieve(src, name[0] + '.png')  # src주소로 이미지를 찾아 '.png'형식으로 저장

    def Page_Screenshot(self):
        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        self.driver.save_screenshot(name[0] + '.png')  # 드라이버가 위치한 페이지에서 스크린샷 후 '.png'형식으로 저장

    def Image_Click(self, ImageFile):
        img = ImageFile  # 이미지 파일을 저장
        pos = imagesearch(img)  # 이미지 파일과 일치하는 이미지의 위치를 반환하는 imagersearch함수호출
        if pos[0] != -1:  # 찾았을 경우
            click_image(img, pos, 'left', 0.1)  # 이미지, 위치, 활성화할 마우스 버튼, 마우스가 원래 위치에서 새 위치로 이동하는 데 걸리는 시간 순

    def Position_Click(self, X, Y):
        pyautogui.click(int(X), int(Y))  # 함수 호출 시 좌표를 찾는 GUI.py -> selectPosition함수 호출
    
    def Delay(self, _time):
        time.sleep(int(_time))  # 지정한 초 동안 딜레이
   
    # 0806 서보인 ImageCrawling 구현중
    def ImageCrawling(self,fordername,filetype):
        
        html = self.driver.page_source # 현재 브라우저의 page_source 다운로드 후 html이라는 변수에 담는다
        soup = BeautifulSoup(html,'html.parser') # beautifulsoup을 이용해 parsing 후 soup 변수에 저장

        n = 1 # 이미지 저장시 고유 번호 

        img = soup.find_all(class_= '_img') # 네이버 이미지 크롤링 시 class 이름 고정 '_img'

        try : 
            for i in img :
                imgurl = i['data-source'] # '_img' class 안에 실제 이미지 경로 = 'data-source'
                with urlopen(imgurl) as f: # urlopen을 통해 이미지 열기
                    with open(fordername +"/" + str(n) + filetype , 'wb')as h : # 사용자 입력값을 받은 경로에 사용자가 지정한 형식으로 이미지 확보 
                        img = f.read() 
                        h.write(img) # 이미지 저장
                n+=1
        except:
            for i in img :
                imgurl = i['src'] # '_img' class 안에 실제 이미지 경로 = 'src'일수도 있음 그래서 두가지 경우의 수로 나누어 크롤링 실시
                with urlopen(imgurl) as f: # urlopen을 통해 이미지 열기
                    with open(fordername +"/" + str(n) + filetype , 'wb')as h : # 사용자 입력값을 받은 경로에 사용자가 지정한 형식으로 이미지 확보 
                        img = f.read() 
                        h.write(img) # 이미지 저장
                n+=1

    # 0806 서보인 resize 구현중
    def Resize(self, Original_Path,Save_Path,X,Y):


        file_list = os.listdir(Original_Path)

        img_list = []
        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        total_image = len(img_list)
        
        index = 1
        result = []
        for name in img_list :

            img = Image.open('%s%s'%(Original_Path+"/", name)) 
            img_array = np.array(img) 
            img_resize = cv2.resize(img_array, (int(X),int(Y)), interpolation = cv2.INTER_AREA) 
            img = Image.fromarray(img_resize) 
            img.save('%s%s'%(Save_Path + "/", name)) 

            result.append(('Resize :' + name + '   ' + str(index) + '/' + str(total_image)))
            index = index + 1
            
     

    # 0807 서보인 ImageSqure 구현중
    def ImageSqure(self,Original_Path,Save_Path):

        files = os.listdir(Original_Path)

        for el in files:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP": 
                image = Image.open(Original_Path + "/" + el) 
                x, y = image.size 
                if x > y :
                    new_size = x 
                    x_offset = 0 
                    y_offset = int((x-y)/2) 
                elif y > x: 
                    new_size = y 
                    x_offset = int((y-x) / 2) 
                    y_offset = 0
            
                background_color = "white"  
                new_image = Image.new("RGB", (new_size, new_size), background_color) 
                new_image.paste(image, (x_offset, y_offset)) 

       
                outfile_name = ".".join(splt) + "." + ext 
                new_image.save(Save_Path + "/" + outfile_name)

    

    # 0807 서보인 ConvertGray 구현중
    def ConvertGray(self,Original_Path,Save_Path):

        file_list = os.listdir(Original_Path)

        img_list = []
        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        total_image = len(img_list)
        
        index = 1
        result = []

        for name in img_list :

            img = Image.open('%s%s'%(Original_Path+"/", name))
            img_array = np.array(img)
            img = Image.fromarray(img_array).convert('L') #위의 이미지 리사이즈하는 기능에서 여기만 바뀜 *상대경로만(str) 있으면 전체 다 회색조로 변환
            img.save('%s%s'%(Save_Path + "/" , name))

       
    # 0807 서보인 Binary 구현중
    def Binary(self,Original_Path,Save_Path):
        orpath = Original_Path
        path = orpath + "/" 
        output_dir =  Save_Path

        file_list = os.listdir(path)

        img_list = []

        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)
        
        total_image = len(img_list)

        imagePaths = [os.path.join(path,file_name) for file_name in os.listdir(path)]
        index = 0

        rgb3 = 255
        rgb4 = 99
        rgb5 = 10

        max_output_value = rgb3   
        neighborhood_size = rgb4   
        subtract_from_mean = rgb5

        result= []

        for name in img_list :

            index = index + 1

            img = Image.open(path+name)
            img_array = np.array(img, 'uint8')

            image_binarized = cv2.adaptiveThreshold(img_array,
                                       max_output_value,
                                       cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                       cv2.THRESH_BINARY,
                                       neighborhood_size,
                                       subtract_from_mean)
            
            cv2.imwrite(output_dir + "/" + name, image_binarized)
    
    # 0807 서보인 EqulizeHistColor 구현중
    def EqulizeHistColor(self,Original_Path,Save_Path):

        orpath = Original_Path
        path = orpath + "/"
        output_dir = Save_Path
        file_list = os.listdir(path)

        img_list = []

        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        total_image = len(img_list)

        index = 0
        result = []

        for name in img_list :
            
            index = index+1
            img = Image.open(path+name)
            img_array = np.array(img)
            img_array = np.array(img, 'uint8')
            image_yuv = cv2.cvtColor(img_array, cv2.COLOR_BGR2YUV)
            image_yuv[:, :, 0] = cv2.equalizeHist(image_yuv[:, :, 0])
            image_rgb = cv2.cvtColor(image_yuv, cv2.COLOR_YUV2RGB)
            cv2.imwrite(output_dir + "/" + name, image_rgb)

       
    # 0807 서보인 EqulizeHistGary 구현중
    def EqulizeHistGray(self,Original_Path,Save_Path):

        orpath = Original_Path
        path = orpath + "/"
        output_dir = Save_Path

        file_list = os.listdir(orpath)
        img_list = []

        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)
        
        total_image = len(img_list)
        index = 0
        result = []

        for name in img_list :
            index = index + 1

            img = Image.open(path+name)
            img_array = np.array(img)

            img_array = np.array(img, 'uint8')
            image_enhanced = cv2.equalizeHist(img_array) 
    
            cv2.imwrite(output_dir + "/" + name, image_enhanced)
    
    # 0807 서보인 GausianBlur 구현중
    def GausianBlur(self,Original_Path,Save_Path,Kernerl_X,Kernerl_Y,Sigma):

        orpath = Original_Path
        path = orpath + "/"
        output_dir = Save_Path

        file_list = os.listdir(path)

        img_list = []

        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        total_image = len(img_list)
        index = 0

        k_width = Kernerl_X 
        k_height = Kernerl_Y  
        sigma = Sigma

        result = []

        for name in img_list :
            index = index + 1

            img = cv2.imread(path+name)
            img_array = np.array(img)

            img_array = np.array(img, 'uint8')
            image_gaussian = cv2.GaussianBlur(img_array, (int(k_width),int(k_height)), int(sigma)) #가우시안 블러 기능
            cv2.imwrite(output_dir + "/" + name, image_gaussian)

        print('GausianBlur가 실행되었습니다.')
    
    # 0807 서보인 MedianBlur 구현중
    def MedianBlur(self,Original_Path,Save_Path,Kernerl_Size):

        orpath = Original_Path
        path = orpath + "/"
        output_dir = Save_Path

        file_list = os.listdir(path)

        img_list = []

        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        total_image = len(img_list)
        index = 0
        Kernel_size = Kernerl_Size

        result = []

        for name in img_list :
            index = index + 1
            img = cv2.imread(path+name)
            img_array = np.array(img, 'uint8')
            image_median = cv2.medianBlur(img_array, int(Kernel_size)) 
            cv2.imwrite(output_dir + "/" + name, image_median)


        print('MedianBlur가 실행되었습니다')
    
    # 0807 서보인 BilateralBlur 구현중
    def BilateralBlur(self,Original_Path,Save_Path,Kernerl_Size,Sigma_Space,Sigma_Color):
        
        orpath = Original_Path
        path = orpath + "/"
        output_dir = Save_Path

        file_list = os.listdir(path)
        img_list = []

        for el in file_list:
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        total_image = len(img_list)
        index = 0

        result = []

        kernerl_size = Kernerl_Size 
        sigma_space = Sigma_Space 
        sigma_color = Sigma_Color

        for name in img_list :
            index = index + 1
            img = cv2.imread(path+name)
            img_array = np.array(img, 'uint8')
            image_bilateral = cv2.bilateralFilter(img_array, int(kernerl_size),int(sigma_space),int(sigma_color)) #쌍방 필터 기능
            cv2.imwrite(output_dir + "/" + name, image_bilateral)
        print('BilateralBlur가 실행되었습니다')

    # 0812 서보인 RowDropNa 구현중
    def RowDropNa(self,filename):
        
        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_RowModify'+'.'+file_full_name[1] # 원본파일명 + 'Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            drop_file_name = read_file_name.dropna() # 선택된 파일의 결측치 제거 dropna 함수를 이용해 결측치가 존재하는 행의 전체를 삭제한다
            drop_file_name.to_excel(file_modify_full_name,index=False) # 원본파일은 그대로 두고 , 결측치가 삭제된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+RowModify.확장자)
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            drop_file_name = read_file_name.dropna() # 선택된 파일의 결측치 제거 dropna 함수를 이용해 결측치가 존재하는 행의 전체를 삭제한다
            drop_file_name.to_csv(file_modify_full_name,index=False) # 원본파일은 그대로 두고 , 결측치가 삭제된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+RowModify.확장자)
    
     # 0812 서보인 ColumnDropNa 구현중
    def ColumnDropNa(self,filename):
        
        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_ColumnModify'+'.'+file_full_name[1] # 원본파일명 + 'Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            drop_file_name = read_file_name.dropna(axis = 1) # 선택된 파일의 결측치 제거 dropna 함수를 이용해 결측치가 존재하는 열의 전체를 삭제한다
            drop_file_name.to_excel(file_modify_full_name,index=False) # 원본파일은 그대로 두고 , 결측치가 삭제된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+ColumnModify.확장자)
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            drop_file_name = read_file_name.dropna(axis = 1) # 선택된 파일의 결측치 제거 dropna 함수를 이용해 결측치가 존재하는 열의 전체를 삭제한다
            drop_file_name.to_csv(file_modify_full_name,index=False) # 원본파일은 그대로 두고 , 결측치가 삭제된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+ColumnModify.확장자)
    
    # 0812 서보인 ColumnDropNa 구현중
    def FillNa(self,filename,text):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_FillNaModify'+'.'+file_full_name[1] # 원본파일명 + 'Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            fill_file_name = read_file_name.fillna(text) # 선택된 파일의 결측치를 사용자 입력값으로 대체한다
            fill_file_name.to_excel(file_modify_full_name,index=False) # 원본파일은 그대로 두고 , 결측치가 사용자 입력값이 삽입된 변경된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+FillNaModify.확장자)
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            fill_file_name = read_file_name.Fillna(text) # 선택된 파일의 결측치를 사용자 입력값으로 대체한다
            fill_file_name.to_csv(file_modify_full_name,index=False) # 원본파일은 그대로 두고 , 결측치가 사용자 입력값이 삽입된 변경된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+FillNaModify.확장자)

    # 0812 서보인 ColumnDropNa 구현중
    def IsNull(self,filename) :

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_IsNullModify'+'.'+file_full_name[1] # 원본파일명 + 'Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            isnull_file_name = read_file_name.isnull().sum() # 선택된 파일의 결측치를 확인한다 
            isnull_file_name.to_excel(file_modify_full_name,header = False) # 원본파일은 그대로 두고 , 결측치를 확인할수 있는 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+IsNullModify.확장자)
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            isnull_file_name = read_file_name.isnull().sum() # 선택된 파일의 결측치를 확인한다
            isnull_file_name.to_csv(file_modify_full_name,header = False) # 원본파일은 그대로 두고 , 결측치를 확인할수 있는 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+IsNullModify.확장자)
    # 0812 서보인 ColumnDropNa 구현중
    def ffill(self,filename):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_ffill_Modify'+'.'+file_full_name[1] # 원본파일명 + 'Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            ffill_file_name = read_file_name.fillna(method='ffill') # 선택된 파일의 결측치를 직전행의 값으로 대체한다
            ffill_file_name.to_excel(file_modify_full_name,index = False) # 원본파일은 그대로 두고 , 결측치가 직전행의 값으로 변경된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+ffill_Modify.확장자)
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            ffill_file_name = read_file_name.fillna(method='ffill') # 선택된 파일의 결측치를 직전행의 값으로 대체한다
            ffill_file_name.to_csv(file_modify_full_name,index = False) # 원본파일은 그대로 두고 , 결측치를 확인할수 있는 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+ffill_Modify.확장자)

        
    # 0812 서보인 ColumnDropNa 구현중
    def bfill(self,filename):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_bfill_Modify'+'.'+file_full_name[1] # 원본파일명 + 'Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            bfill_file_name = read_file_name.fillna(method='bfill') # 선택된 파일의 결측치를 직전행의 값으로 대체한다
            bfill_file_name.to_excel(file_modify_full_name,index = False) # 원본파일은 그대로 두고 , 결측치가 직전행의 값으로 변경된 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+bfill_Modify.확장자)
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            bfill_file_name = read_file_name.fillna(method='bfill') # 선택된 파일의 결측치를 직전행의 값으로 대체한다
            bfill_file_name.to_csv(file_modify_full_name,index = False) # 원본파일은 그대로 두고 , 결측치를 확인할수 있는 파일을 새로 생성한다.(생성되는 파일의 이름 형식은 : 원본파일+bfill_Modify.확장자)
    # 0814 서보인 ColumnDropNa 구현중
    def mean(self,filename):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_mean_Modify'+'.'+file_full_name[1] # 원본파일명 + 'mean_Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            column_count = list(read_file_name) # 파일의 최상위 항목을 리스트 형태로 가져온다 
            for i in column_count: # 리스트 형태로 만든 열의 갯수만큼 반복문 실행
                mean_value = read_file_name[i].mean() # 각 열의 평균값을 mena_value에 담는다
                read_file_name[i] = read_file_name[i].fillna(mean_value) # 각 열의 결측치에 평균값으로 대체한다
            read_file_name.to_excel(file_modify_full_name,index = False) # 반복문이 완료되고 나서 결측치가 평균값으로 대체된 파일이 새로 생성된다.
        
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            column_count = list(read_file_name) # 파일의 최상위 항목을 리스트 형태로 가져온다 
            for i in column_count: # 리스트 형태로 만든 열의 갯수만큼 반복문 실행
                mean_value = read_file_name[i].mean() # 각 열의 평균값을 mena_value에 담는다
                read_file_name[i] = read_file_name[i].fillna(mean_value) # 각 열의 결측치에 평균값으로 대체한다
            read_file_name.to_csv(file_modify_full_name,index = False) # 반복문이 완료되고 나서 결측치가 평균값으로 대체된 파일이 새로 생성된다.
    # 0814 서보인 ColumnDropNa 구현중
    def median(self,filename):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_median_Modify'+'.'+file_full_name[1] # 원본파일명 + 'mean_Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다

        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            column_count = list(read_file_name) # 파일의 최상위 항목을 리스트 형태로 가져온다 
            for i in column_count: # 리스트 형태로 만든 열의 갯수만큼 반복문 실행
                median_value = read_file_name[i].median() # 각 열의 중앙값을 median_value에 담는다
                read_file_name[i] = read_file_name[i].fillna(median_value) # 각 열의 결측치에 중앙값으로 대체한다
            read_file_name.to_excel(file_modify_full_name,index = False) # 반복문이 완료되고 나서 결측치가 중앙값으로 대체된 파일이 새로 생성된다.
        
        else:
            read_file_name = pd.read_csv(filename) # 선택된 파일을 pandas로 읽어온다
            column_count = list(read_file_name) # 파일의 최상위 항목을 리스트 형태로 가져온다 
            for i in column_count: # 리스트 형태로 만든 열의 갯수만큼 반복문 실행
                median_value = read_file_name[i].median() # 각 열의 중앙값을 median_value에 담는다
                read_file_name[i] = read_file_name[i].fillna(median_value) # 각 열의 결측치에 중앙값으로 대체한다
            read_file_name.to_csv(file_modify_full_name,index = False) # 반복문이 완료되고 나서 결측치가 중앙값으로 대체된 파일이 새로 생성된다.
   
    
    # 0819 서보인 mode 구현중
    def mode(self,filename):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_mode_Modify'+'.'+file_full_name[1] # 원본파일명 + 'mode_Modify' + '.' + 확장자 형식으로 저장할 파일명을 만든다

        if(file_full_name[1] == 'xlsx') : # 선택된 파일의 확장자를 비교하여 excel , csv 둘다 파일을 읽을수 있도록 처리한다
            read_file_name = pd.read_excel(filename) # 선택된 파일을 pandas로 읽어온다
            column_count = list(read_file_name) # 파일의 최상위 항목을 리스트 형태로 가져온다 
            for i in column_count: # 리스트 형태로 만든 열의 갯수만큼 반복문 실행
                numbers = read_file_name[i] # 리스트 형태로 빈도수가 높은 값을 찾기위한 변수를 만든다
                frequent = self.modefinder(numbers) # 위에 선언한 modefinder 함수를 통해 빈도수가 가장 높은 값을 찾는다 
                frequents = frequent[0] # modefinter 함수의 반환값 형식은 list 이기에 , 실수형 값으로 변환한다.
                read_file_name[i] = read_file_name[i].fillna(frequents) # 각 열의 결측치에 최빈값으로 대체한다
            read_file_name.to_excel(file_modify_full_name,index = False) # 반복문이 완료되고 나서 결측치가 최빈값으로 대체된 파일이 새로 생성된다.
    # 0819 서보인 kor_ocr 구현중
    def Kor_Ocr(self,fordername):

        pytesseract.pytesseract.tesseract_cmd = r'C:\Users\PC\Desktop\RPA\RPA\Tesseract-OCR\tesseract' # 경로는 사용자의 경로에 맞게 재설정
        orpath = fordername
        path = orpath + "/"

        file_list = os.listdir(path)
        img_list = []
        result = []

        for el in file_list: # 파일 타입이 이미지 형식인지 확인하는 반복문 
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        for name in img_list : # 실질적으로 폴더 안의 이미지 텍스트 추출 하는 반복문

            img = Image.open(path+name)
            img_array = np.array(img)
            
            addim = pytesseract.image_to_string(img_array,lang='kor')
            result.append(addim)
            print(addim)
        
        result_data = pd.DataFrame(result)
        result_data.to_excel(fordername+'_Ocr_Kor'+'.'+'xlsx',index = False)

        print('Kor_Ocr이 실행되었습니다.')
    # 0819 서보인 eng_ocr 구현중
    def Eng_Ocr(self,fordername):
        
        pytesseract.pytesseract.tesseract_cmd = r'C:\Users\PC\Desktop\RPA\RPA\Tesseract-OCR\tesseract' # 경로는 사용자의 경로에 맞게 재설정
        orpath = fordername
        path = orpath + "/"

        file_list = os.listdir(path)
        img_list = []
        result = []

        for el in file_list: # 파일 타입이 이미지 형식인지 확인하는 반복문 
            splt = el.split(".")
            ext = splt.pop()
            if ext in "jpg jpeg png bmp JPG JPEG PNG BMP":
                img_list.append(el)

        for name in img_list : # 실질적으로 폴더 안의 이미지 텍스트 추출 하는 반복문

            img = Image.open(path+name)
            img_array = np.array(img)
            
            addim = pytesseract.image_to_string(img_array,lang='Eng')
            result.append(addim)
            print(addim)
        
        result_data = pd.DataFrame(result)
        result_data.to_excel(fordername+'_Ocr_Eng'+'.'+'xlsx',index = False)


        print('Eng_Ocr이 실행되었습니다.')


    # 0821 서보인 Filter_Content 구현중
    def Filter_Content(self,filename):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_Filter_Content'+'.'+file_full_name[1] # 원본파일명 + '_Filter_Content' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        if file_full_name[1] == 'xlsx' :
            filter_content = [] # 파일의 불용어 처리를 끝낸 후 데이터를 저장할 변수 
            filename_pandas = pd.read_excel(filename) # 판다스로 데이터 읽기
            pandas_tolist = filename_pandas.values.tolist() # 판다스로 읽어온 데이터 list로 변환 
            for i in pandas_tolist : # 리스트를 반복문을 통해 실질적인 불용어 처리 
                filter_content.append((str(i).replace('.', '').replace(',','')
                .replace("'","").replace('·', ' ').replace('=','').replace('\\n','').replace('[','')
                .replace(']','').replace('-','').replace('"','').replace('”','').replace('|','')
                .replace('||','').replace('/','').replace('`','').replace('「','').replace(';','')
                .replace('?','').replace('!','').replace(':','').replace('>','').replace('<','')
                .replace('@','').replace('」',''))) 
            result = pd.DataFrame(filter_content) # 다시 리스트를 판다스로 데이터 형식을 변환
            result.to_excel(file_modify_full_name,index=False) # 판다스를 이용한 데이터 xlsx 저장
        else : 
            filter_content = [] # 파일의 불용어 처리를 끝낸 후 데이터를 저장할 변수 
            filename_pandas = pd.read_excel(filename) # 판다스로 데이터 읽기
            pandas_tolist = filename_pandas.values.tolist() # 판다스로 읽어온 데이터 list로 변환 
            for i in pandas_tolist : # 리스트를 반복문을 통해 실질적인 불용어 처리 
                filter_content.append((str(i).replace('.', '').replace(',','')
                .replace("'","").replace('·', ' ').replace('=','').replace('\\n','').replace('[','')
                .replace(']','').replace('-','').replace('"','').replace('”','').replace('|','')
                .replace('||','').replace('/','').replace('`','').replace('「','').replace(';','')
                .replace('?','').replace('!','').replace(':','').replace('>','').replace('<','')
                .replace('@','').replace('」',''))) 
            result = pd.DataFrame(filter_content) # 다시 리스트를 판다스로 데이터 형식을 변환
            result.to_csv(file_modify_full_name,index=False ) # 판다스를 이용한 데이터 csv 저장
    
    # 0828 서보인 Morphs
    def Morphs(self, filename):
        
        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_Morphs'+'.'+file_full_name[1] # 원본파일명 + '_Filter_Content' + '.' + 확장자 형식으로 저장할 파일명을 만든다

        Okt = konlpy.tag.Okt() # 품사태깅을 위한 라이브러리 객체 생성
        Okt_morphs = [] # 품사태깅이 완료되고 나서 모든 데이터 값을 저장할 변수 생성
        
        if file_full_name[1] == 'xlsx' :
            filename_pandas = pd.read_excel(filename) # 사용자 입력값으로 받은 파일 경로를 판다스로 읽어옴
            pandas_tolist = filename_pandas.values.tolist() # 반복문을 사용하기 위해 데이터 타입을 pandas에서 list로 변환
            for i in pandas_tolist : # 실질적인 품사태깅이 이루어지는 곳
                for ii in i: 
                    Okt_morphs.append(Okt.pos(ii)) # 두번의 반복문을 통해 하나의 셀안에 1개의 값을 입력
            list_topandas = pd.DataFrame(Okt_morphs) # Pandas를 이용해 액셀을 저장하기 위해 데이터 타입 pandas에서 list로 변환
            list_topandas.to_excel(file_modify_full_name,index=False) # 액셀파일로 결과물 저장
        
        else : 
            filename_pandas = pd.read_csv(filename) # 사용자 입력값으로 받은 파일 경로를 판다스로 읽어옴
            pandas_tolist = filename_pandas.values.tolist() # 반복문을 사용하기 위해 데이터 타입을 pandas에서 list로 변환
            for i in pandas_tolist : # 실질적인 품사태깅이 이루어지는 곳
                for ii in i: 
                    Okt_morphs.append(Okt.pos(ii)) # 두번의 반복문을 통해 하나의 셀안에 1개의 값을 입력
            list_topandas = pd.DataFrame(Okt_morphs) # Pandas를 이용해 csv을 저장하기 위해 데이터 타입 pandas에서 list로 변환
            list_topandas.to_csv(file_modify_full_name,index=False) # csv파일로 결과물 저장
    
    # 0828 서보인 완료
    def Extract_Word(self,filename):
        
        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_Extract_Word'+'.'+file_full_name[1] # 원본파일명 + '_Filter_Content' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        filename_pandas = pd.read_excel(filename) # 판다스로 액셀파일을 불러옴
        pandas_tolist = filename_pandas.values.tolist() # 판다스 형태의 데이터 타입을 리스트 형식으로 변환
        Noun_words = [] # 액셀파일로 부터 불러온 데이터의 전처리를 위한 변수
        Noun_words_list = [] # 전처리가 끝난 데이터를 새로운 리스트에 저장
        tmp = [] # 임시 변수를 이용해 각 행마다의 전처리된 데이터를 담을 임시변수를 생성
        

        for i in pandas_tolist : # 잘못된 방법임 하지만 결과물은0 나옴
            for ii in i :
                test1 = (str(ii).replace(",",""))
                test2 = test1.replace("(","").replace(")","")
                test3 = test2.replace("'","")
                test4 = test3.replace(" ",",")
                test5 = test4.split(",")
                tmp.append(test5) # 반복문을 통해 얻어온 데이터를 임시변수에 담아 각 행에 가공된 데이터를 담는다 (불용어 제거  ex) '(' , ')' , "'" , "'")
            Noun_words.append(tmp) # 각 행마다의 전처리된 데이터를 3차원 배열로 생성 
            tmp = [] 

        for i in Noun_words : # 명사를 추출하기 위한 반복문
            for ii in i :
              
                try :
                    if ii[1] == 'Noun':
                        tmp.append(ii[0])
                except:
                    pass
            Noun_words_list.append(tmp) # 전체 데이터를 한곳에 모은다
            tmp = [] # 각 배열에 대한 반복문이 끝날시 임시변수를 초기화 한다

        result = pd.DataFrame(Noun_words_list) # 리스트 형태의 데이터 판다스 타입으로 변환 
        result.to_excel(file_modify_full_name,index = False) # 데이터를 액셀파일로 저장
   
    # 붙이기 실패
    def plot_bar(self,filename):

        file_name = filename # 원본파일 이름 불러오기
        file_full_name = file_name.split('.') # split함수를 통한 파일명과 확장자 분리
        file_first_name = file_full_name[0] # 원본파일 이름
        file_second_name = file_full_name[1] # 확장자 이름

        if(file_second_name == 'xlsx') : # xlsx 와 , csv 파일 분리
            read_file_name = pd.read_excel(file_name) # 확장자에 따른 판다스 데이터 읽기
            read_file_name.plot.bar() # 막대 그래프로 그래프를 표현 하겠다.
            plt.title(filename) # 그래프 이름 설정
            plt.xlabel('test1') # X축
            plt.ylabel('test2') # Y축
            plt.show() # 그래프를 화면상에 보임
            plt.legend(['test1','test2','test3','test4']) # 범례 표현
            print(read_file_name.head())
            print('xlsx 파일 입니다.')

        else :
            print('csv 파일 입니다.')

    # 붙이기 실패
    def plot_line(self,filename):

        file_name = filename # 원본파일 이름 불러오기
        file_full_name = file_name.split('.') # split함수를 통한 파일명과 확장자 분리
        file_first_name = file_full_name[0] # 원본파일 이름
        file_second_name = file_full_name[1] # 확장자 이름

        if(file_second_name == 'xlsx') : # xlsx 와 , csv 파일 분리
            read_file_name = pd.read_excel(file_name) # 확장자에 따른 판다스 데이터 읽기
            read_file_name.plot() # 꺽은선으로 그래프를 표현 하겠다.
            plt.title(filename) # 그래프 이름 설정
            plt.xlabel('test1') # X축
            plt.ylabel('test2') # Y축
            plt.show() # 그래프를 화면상에 보임
            plt.legend(['test1','test2','test3','test4']) # 범례 표현
            print(read_file_name.head())
            print('xlsx 파일 입니다.')

        else :
            print('csv 파일 입니다.')

    def Word_Cloud(self,filename):

        file_name = filename # 저장할 파일의 이름의 원본을 가져온다
        file_full_name = file_name.split('.') # 원본파일에 변경됨을 알리기 위한 Modify 문자열을 삽입하기 위해 split 함수를 통해 파일경로명과 , 확장자를 분리한다
        file_modify_full_name = file_full_name[0]+'_Word_Cloud'+'.'+file_full_name[1] # 원본파일명 + '_Word_Clout' + '.' + 확장자 형식으로 저장할 파일명을 만든다
        
        filename_pandas = pd.read_excel(filename).fillna('') # 판다스로 액셀파일을 불러옴
        pandas_tolist = filename_pandas.values.tolist() # 판다스 형태의 데이터 타입을 리스트 형식으로 변환 , 2차원 리스트
        result = [] # wordcloud에 사용될 변수
        
        for i in pandas_tolist : # 1차원 배열로 재생성
            for ii in i :
                if ii != '': # 위 판다스로 데이터를 불러올 때 빈 공백을 nan 값으로 인식함 , 그걸 '' 공백값으로 바꿔주었는데 빈 공백값을 제외한 명사만 
                    result.append(ii) # result라는 변수에 담는 작업

        count = Counter(result) # 1차원 배열의 대한 단어 count 리스트 생성
        list_count = (count.most_common()) # 빈도수 파악 
        dict_count = dict(list_count) # 리스트를 딕셔너리로 생성
        

        FONT_PATH = 'C:/Windows/Fonts/malgun.ttf' # For Korean characters
        wordcloud = WordCloud(max_font_size=60, relative_scaling=.5, font_path=FONT_PATH).generate_from_frequencies(dict_count) # 워드클라우드 옵션들 
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.show() # 그래프 보여주기
        


        

        
        


                
           
            
            










        