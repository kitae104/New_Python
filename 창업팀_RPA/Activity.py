# -*- coding: utf-8 -*-

from PyQt5.QtWidgets import QFileDialog
from selenium import webdriver
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from openpyxl import load_workbook
from email import encoders
from bs4 import BeautifulSoup
import smtplib
import urllib.request
import openpyxl
import xlrd
import operator
import docx2txt
from docx import Document
from PyQt5.QtWidgets import *
import time
import urllib.request
from imagesearch import *
import pyautogui
import sys
import re
from docx.shared import Cm
import csv

class Activity:

    driver = webdriver
    id = ''
    class_name = ''
    sid = ''
    part = ''
    attachfile = 0
    wb = ''
    ws = ''
    Name = ''
    list = ''
    document = ''
    rdr = ''
    wr = ''
    f = ''

    def __init__(self):
        self.__activity_dict = {  # 기능들의 이름, 변수, 변수값을 기본값으로 저장
            '브라우저 열기': {
                'url': '_url'
            },
            'id값으로 클릭': {
                'id': '_id'
            },
            'id값으로 입력': {
                'id': '_id', 'key': '_key'
            },
            'class값으로 클릭': {
                'class_name': '_class_name'
            },
            'class값으로 입력': {
                'Class': '_Class', 'key': '_key'
            },
            'css값으로 클릭': {
                'Css': '_Css'
            },
            'css값으로 입력': {
                'Css': '_Css', 'key': '_key'
            },
            'xpath값으로 클릭': {
                'XPath': '_XPath'
            },
            'xpath값으로 입력': {
                'XPath': '_XPath', 'key': '_key'
            },
            'Gmail 세션 생성': {
                'id': '_id', 'smtppw': '_smtppw'
            },
            '파일첨부': {
                'filename': '_filename'
            },
            '엑셀 읽어 보내기': {
                'filename': '_filename'
            },
            'text 크롤링': {
                'class_name': 'class_name'
            },
            'url 크롤링': {
                'class_name': 'class_name', 'company_class': '_company_class'
            },
            '텍스트 파일 읽기': {
                'filename': '_filename'
            },
            '텍스트 파일 쓰기': {
                'text': '_text'
            },
            '텍스트 추가': {
                'filename': '_filename', 'text': '_text'
            },
            '텍스트 파일을 엑셀 파일로 변환': {
                'text': '_text'
            },
            '시트 불러오기': {
                'filename': '_filename'
            },
            '행 삭제': {
                'row': '_row'
            },
            '열 삭제': {
                'col': '_col'
            },
            '셀 옮기기': {
                'cell': '_cell', 'move': '_move'
            },
            '값 얻기': {
                'cell': '_cell'
            },
            '행 찾기': {
                'row': '_row', 'text': '_text', 'totalrow': '_totalrow'
            },
            '값 설정': {
                'cell': '_cell', 'text': '_text'
            },
            'docx 읽기': {
                'filename': '_filename'
            },
            'docx 생성': {
                'style': '_style'
            },
            '이미지 추가': {
                'filename': '_filename', 'size': '_size'
            },
            'docx 작성': {
                'style': '_style'
            },
            'head 쓰기': {
                'text': '_text', 'style': '_style'
            },
            'body 쓰기': {
                'text': '_text', 'style': '_style'
            },
            'docx 저장': {

            },
            'csv 읽기': {
                'filename': '_filename'
            },
            '열 가져오기': {
                'selline': '_selline'
            },
            'csv 쓰기': {
                'text': '_text'
            },
            '행 추가': {
                'text': '_text'
            },
            'csv파일 추가모드로 열기': {
                'filename': '_filename'
            },
            '파일 닫기': {

            },
            '모든 값 얻기': {

            },
            '브라우저 닫기': {

            },
            '팝업창 닫기': {

            },
            '이미지 다운로드(src)': {
                'src': '_src'
            },
            '페이지 스크린샷': {

            },
            '이미지 클릭': {
                'ImageFile': '_imagefilename'
            },
            'url 여러 페이지 크롤링': {
                'class_name': '_class_name', 'company_class': '_company_class', 'count': '_count'
            },
            '데이터로 크롤링': {

            },
            '좌표 클릭': {
                'X': 'Pos_X', 'Y': 'Pos_Y'
            },
            '딜레이': {
                'time': '_time'
            },
            '반복_For': {
                'startnum': '_startnum', 'endnum': '_endnum', 'roofnum': '_roofnum'
            }

        }

    # key의 갯수 세기
    def count_keys(self, name):
        keylist = list(self.__activity_dict[name].keys())  # key값들 저장
        return len(keylist)

    # 번호에 해당하는 key의 값
    def search_name_key(self, name, num):
        keylist = list(self.__activity_dict[name].keys())   # key값들 저장
        return keylist[num]

    # key에 해당하는 value
    def search_key_value(self, name, key):
        value = self.__activity_dict[name][key]  # key에 해당하는 value
        return value

    def OpenBrowser(self, url):
        driver = webdriver.Chrome()  # 크롬 웹 드라이버 생성
        driver.maximize_window()  # 페이지창 최대화
        driver.implicitly_wait(3)  # 3초대기
        driver.get(url)  # url주소로 이동
        self.driver = driver  # 드라이버 전역변수로 저장

    def find_element_by_id_click(self, id):
        self.id = self.driver.find_element_by_id(id)
        self.id.click()

    def find_element_by_id_sendkey(self, id, key):
        self.driver.find_element_by_id(id).send_keys(key)

    def find_element_by_class_click(self, _class):
        self.class_name = self.driver.find_element_by_class_name(_class)
        self.class_name.click()

    def find_element_by_class_sendkey(self, _class, key):
        self.driver.find_element_by_class_name(_class).send_keys(key)

    def find_element_by_css_click(self, Css):
        self.driver.find_element_by_css_selector(Css).click()

    def find_element_by_css_sendkey(self, Css, key):
        self.driver.find_element_by_css_selector(Css).send_keys(key)

    def find_element_by_xpath_click(self, XPath):
        self.driver.find_element_by_xpath(XPath).click()

    def find_element_by_xpath_sendkey(self, XPath, key):
        self.driver.find_element_by_xpath(XPath).send_keys(key)

    def Gmail_Session_Create(self, id, pw):  # STMP세션 생성(GMAIL)
        self.session = smtplib.SMTP("smtp.gmail.com", 587)  # 세션생성
        self.session.starttls()  # SMTP 연결을 TLS (전송 계층 보안) 모드로 전환, 그뒤에 모든 STMP명령은 암호화되기 위함
        self.sid = id  # STMP id저장
        self.session.login(id, pw)  # SMTP 계정 인증

    def Send_Mail(self, name, receiver, subject, content):  # 메일 전송

        self.msg = MIMEMultipart()  # MIME 메세지 생성

        self.msg["From"] = self.sid  # 메세지의 송신자 설정
        self.msg["To"] = receiver  # 메세지의 수신자 저장
        self.msg["Subject"] = subject  # 메세지의 내용 저장

        text = MIMEText(_text=content)
        self.msg.attach(text)  # 내용 첨부
        if self.attachfile == 1:  # 파일을 첨부할 경우 : attachfile == 1
            self.msg.attach(self.part)  # part에 저장된 파일 첨부
        self.session.sendmail(self.sid, receiver, self.msg.as_string())  # 메일 전송(STMPid, 메세지의 수신자, 메세지의 내용)

    def Attach_File(self, Filename):  # 파일 첨부
        path = Filename  # 파일이름 저장
        with open(path, 'rb') as f:  # 파일 열기
            self.part = MIMEBase("application", "octet-stream")  # 이진 파일 형식으로 변환
            self.part.set_payload(f.read())  # 7계층 프로토콜 적용
            encoders.encode_base64(self.part)  # 영상, 이미지 파일을 문자열 형태로 변환
            self.part.add_header('Content-Disposition', 'attachment', filename=path)  # Content-Disposition헤더를 attachment속성으로 명시하여 파일 첨부를 활성화
            self.attachfile = True  # 파일이 있음을 의미

    def Excel_read_send(self, excelFile):  # 엑셀 파일 읽은 후 전송 함수 호출
        wb = load_workbook(excelFile)  # 엑셀파일 불러와 워크북 생성
        ws = wb.active  # 워크시트 생성
        for row in ws.iter_rows():  # 열단위로 읽어 name, receiver, subject, content에 저장
            name = row[1].value
            receiver = row[2].value
            subject = row[3].value
            content = row[4].value
            Activity.Send_Mail(self, name, receiver, subject, content)  # 전송 함수 호출
        self.session.quit()  # 세션 종료

    def WebCrawling_text(self, _class):  # 크롤링해서 text를 읽고 txt파일로 저장
        # 파일 지정해서 열기
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        f = open(file_name, "w", encoding="utf-8")

        # 크롤링
        h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함 (서버 측에서 자동 프로그램이 동작하지 않게 하기위해 막아 놓은거 푸는 기능)
        url = self.driver.current_url  # 브라우저의 현재 url을 가져온다.
        req = urllib.request.Request(url, headers=h)  # 벤 걸리지 않게 함
        reqhttp = urllib.request.urlopen(req)
        http = reqhttp.read()
        soup = BeautifulSoup(http, 'html.parser')

        class_choice = soup.find_all(class_=_class)  # 크롤링할 class에서 가져온 값을 class_choice에 저장

        for count in class_choice:
            text = str(count.get_text().strip()) + "\n"  # 텍스트를 단위별로 읽어와서 text에 태그들 제거하고 저장한다.
        text = text.replace("// flash 오류를 우회하기 위한 함수 추가", "")
        text = text.replace("function _flash_removeCallback() {}", "")
        f.write(text)  # text에 저장된 내용을 file에 쓴다.
        f.close()

    def WebCrawling_url(self, _class, _company_class):  # 크롤링해서 text를 읽고 text에 저장된 하이퍼링크를 가져와서 엑셀 파일로 저장
        wb = openpyxl.Workbook()  # 워크북 생성
        ws = wb.active  # 워크시트 생성

        # 크롤링
        h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함
        url = self.driver.current_url
        print(url)
        req = urllib.request.Request(url, headers=h)  # 벤 걸리지 않게 함
        reqhttp = urllib.request.urlopen(req)
        http = reqhttp.read()  # .decode('utf-8') #읽고 한글 깨지지 않게 함
        soup = BeautifulSoup(http, 'html.parser')

        class_choice = soup.find_all(class_=_class)  # 크롤링할 class에서 가져온 값을 class_choice에 저장
        data = []
        text = []
        company = []
        for n in class_choice:
            textsave = str(n.get_text().strip())
            datasave = n["href"]  # 크롤링해서 가져온 정보를 href로 변환해서 datasave에 저장
            data.append(datasave)  # datasave에 저장된 값을 순서대로 data리스트에 저장
            text.append(textsave)  # textsave에 저장된 값을 순서대로 text리스트에 저장

        company_class_choice = soup.find_all(class_=_company_class)
        for n in company_class_choice:
            datasave = str(n.get_text().strip())
            datasave = datasave.replace("언론사 선정", "")
            company.append(datasave)

        for data_range in range(len(data)):  # data에 저장된 개수만큼 반복문을 돌림
            ws.cell(row=data_range + 1, column=1).value = data[data_range]  # 엑셀 파일에 순서대로 저장 (1,1), (2,1) ....
            ws.cell(row=data_range + 1, column=2).value = text[data_range]
            ws.cell(row=data_range + 1, column=3).value = company[data_range]

        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        wb.save(name[0] + '.xlsx')

    def WebCrawling_url_count(self, _class, _company_class, _count):  # 페이지 수를 지정하여 크롤링함.(WebCrawling_url과 페이지 수 지정을 제외하면 동일)
        wb = openpyxl.Workbook()  # 워크북 생성
        ws = wb.active  # 워크시트 생성

        # 크롤링
        h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함
        first_url = self.driver.current_url  # 현재 url을 가져옴
        req = urllib.request.Request(first_url, headers=h)  # 벤 걸리지 않게 함
        reqhttp = urllib.request.urlopen(req)
        http = reqhttp.read()  # .decode('utf-8') #읽고 한글 깨지지 않게 함
        soup = BeautifulSoup(http, 'html.parser')

        # 페이지 수 지정
        pageNum = 1
        lastPage = int(_count) * 10 - 8
        data = []  # 엑셀에 저장되어 있는 url을 저장할 변수 생성
        text = []  # 엑셀에 저장되어 있는 제목을 저장할 변수 생성
        company = []  # 엑셀에 저장되어 있는 회사 이름을 저장할 변수 생성

        while pageNum < lastPage:  # 지정된 페이지 수 보다 작을 경우 반복문 실행
            pageCheck = "&start=" + str(pageNum) + "&refresh_start=0"  # 페이지 수를 지정할 내용
            url = first_url + pageCheck  # 가져온 url뒤에 페이지 수에 대한 정보를 이어 붙임
            html = urllib.request.urlopen(url).read()
            soup = BeautifulSoup(html, 'html.parser')

            class_choice = soup.find_all(class_=_class)  # 신문기사 제목의 클래스 이름(제목만 크롤링하면, 제목과 url을 가져올 수 있음)
            for n in class_choice:
                textsave = str(n.get_text().strip())
                datasave = n["href"]  # 크롤링해서 가져온 정보를 href로 변환해서 datasave에 저장
                data.append(datasave)  # datasave에 저장된 값을 순서대로 data에 저장
                text.append(textsave)  # textsave에 저장된 값을 순서대로  text에 저장

            company_class_choice = soup.find_all(class_=_company_class) # 신문기사 회사의 클래스 이름
            for n in company_class_choice:
                datasave = str(n.get_text().strip()) # 회사의 이름을 읽음
                datasave = datasave.replace("언론사 선정", "")  # 네이버에서 언론사 선정이라는 값을 추가로 입력할 때가 있기 때문에, 회사이름만 크롤링하기 위해 "언론사 선정" 제거
                company.append(datasave)  # datasave에 저장된 값을 company에 저장

            pageNum += 10  # 네이버에서 다음 페이지로 넘어갈 때, 필요한 값

        for data_range in range(len(data)):  # data, text, company에 저장된 값을 엑셀에 저장
            ws.cell(row=data_range + 1, column=1).value = data[data_range]
            ws.cell(row=data_range + 1, column=2).value = text[data_range]
            ws.cell(row=data_range + 1, column=3).value = company[data_range]

        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        wb.save(name[0] + '.xlsx')

    def WebCrawling_get_url(self):  # WebCrawling_url, WebCrawling_url_count를 사용하여 생성된 엑셀파일을 읽어와서 그 값으로 크롤링하는 함수
        name = QFileDialog.getOpenFileName()  # 파일 열기
        wb = xlrd.open_workbook(name[0])  # 엑셀파일 열기
        ws = wb.sheet_by_index(0)  # 첫번째 시트
        numrow = ws.nrows  # 열의 수 numrow에 저장

        data_url = []
        for row in range(numrow):
            if (ws.row_values(row)[0]):  # row에 값이 있다면(url)
                text = ws.row_values(row)[0]  # 그 값을 text에 저장
                data_url.append(text)  # text에 저장된 값을 data_url에 저장

        data_text = []
        for row in range(numrow):
            if (ws.row_values(row)[1]):  # row에 값이 있다면(제목)
                text = ws.row_values(row)[1]  # 그 값을 text에 저장
                data_text.append(text)  # text에 저장된 값을 data_text에 저장

        data_company = []
        for row in range(numrow):
            if (ws.row_values(row)[2]):  # row에 값이 있다면(회사 이름)
                text = ws.row_values(row)[2]  # 그 값을 text에 저장
                data_company.append(text)  # text에 저장된 값을 data_company에 저장

        dirname = QFileDialog.getExistingDirectory()  # 크롤링한 txt파일을 저장할 폴더 지정

        for num in range(len(data_url)):
            file_name = data_text[num] + ".txt"  # 신문기사 제목으로 txt파일 이름 지정
            file_name = re.sub('[-=+,#/\?:^$@*\"※~&%ㆍ!』"\\‘|\(\)\[\]\<\>`\'》]', '', file_name)  # 제목에 특수문자가 존재한다면 제거
            file_name = dirname + '/' + file_name  # txt파일 저장할 폴더와 이름 지정
            f = open(file_name, "w", encoding="utf-8")
            # 크롤링
            h = {'User-Agent': 'Mozilla/5.0'}  # 벤 걸리지 않게 함 (서버 측에서 자동 프로그램이 동작하지 않게 하기위해 막아 놓은거 푸는 기능)
            url = data_url[num]
            req = urllib.request.Request(url, headers=h)  # 벤 걸리지 않게 함
            reqhttp = urllib.request.urlopen(req)
            http = reqhttp.read()
            soup = BeautifulSoup(http, 'html.parser')

            # 회사별로 크롤링
            if data_company[num] == '연합뉴스':
                class_choice = soup.find_all(class_='story-news article')
            elif data_company[num] == '파이낸셜뉴스':
                class_choice = soup.find_all(class_='article_cont mg cont_txt_read')
            elif data_company[num] == '뉴시스':
                class_choice = soup.find_all(id_='textBody')
            elif data_company[num] == '스포츠서울':
                class_choice = soup.find_all(class_='news_text')
            elif data_company[num] == '뉴스핌':
                class_choice = soup.find_all(id_='news_contents')
            elif data_company[num] == '헬스조선':
                class_choice = soup.find_all(class_='par')
            elif data_company[num] == '머니투데이':
                class_choice = soup.find_all(class_='view_content')
            elif data_company[num] == '한국경제TV':
                class_choice = soup.find_all(class_='box-news-body')
            elif data_company[num] == '서울경제':
                class_choice = soup.find_all(class_='article_view')
            elif data_company[num] == 'SBS':
                class_choice = soup.find_all(class_='text_area')
            elif data_company[num] == '조세일보':
                class_choice = soup.find_all(articleBody_='articleBody')
            elif data_company[num] == '오마이뉴스':
                class_choice = soup.find_all(class_='at_contents')
            elif data_company[num] == '뉴스1':
                class_choice = soup.find_all(class_='detail sa_area')
            elif data_company[num] == '노컷뉴스':
                class_choice = soup.find_all(class_='articleBody')
            elif data_company[num] == '한국경제':
                class_choice = soup.find_all(id_='articletxt')
            elif data_company[num] == 'KBS':
                class_choice = soup.find_all(class_='detail-body font-size')
            elif data_company[num] == '스포츠동아':
                class_choice = soup.find_all(class_='article_word')
            elif data_company[num] == 'ZDNet Korea':
                class_choice = soup.find_all(class_='view_cont')
            elif data_company[num] == 'IT조선':
                class_choice = soup.find_all(class_='article font03')
            elif data_company[num] == '동아일보':
                class_choice = soup.find_all(class_='article_txt')
            elif data_company[num] == '국민일보':
                class_choice = soup.find_all(class_='tx')
            elif data_company[num] == '블로터':
                class_choice = soup.find_all(class_='article--content')
            elif data_company[num] == 'MBC':
                class_choice = soup.find_all(class_='news_txt')
            elif data_company[num] == '헤럴드경제':
                class_choice = soup.find_all(class_='con_left_bd')
            elif data_company[num] == '경향신문':
                class_choice = soup.find_all(class_='art_body')
            else:
                class_choice = soup.find_all(class_='')

            for class_count in class_choice:
                text = str(class_count.get_text().strip()) + "\n"  # 텍스트를 단위별로 읽어와서 text에 태그들 제거하고 저장한다.
            text = text.replace("// flash 오류를 우회하기 위한 함수 추가", "")
            text = text.replace("function _flash_removeCallback() {}", "")
            f.write(text)  # text에 저장된 내용을 file에 쓴다.
            f.close()

    def Read_TextFile(self, fName):  # 파일이름으로 txt 파일 열기
        f = open(fName, "r", encoding="utf8")  # 텍스트 파일 read 형태로 열기
        self.list = f.readlines()  # list에 txt파일 모든 줄 저장하기

    def Write_TextFile(self, text):  # 새로운 txt파일쓰기
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"  # 다른이름으로 파일 저장
        f = open(file_name, "w+t", encoding="utf-8")  # 위에서 저장한 파일 쓰기 형태로 열기
        f.write(text+'\n')  # text에 입력된 값을 txt에 쓴다.

    def Add_Text(self, filename, text):  # 기존 txt에 추가하기
        f = open(filename, "at", encoding='utf-8')  # txt파일을 수정 형태로 연다.
        f.writelines(text+'\n')  # 입력된 값을 한줄씩 입력한다.

    def txt_to_xlsx(self, Division):  # txt 내에서 데이터를 구별할 값 입력
        wb = openpyxl.Workbook()  # sheet 만들기
        sheet = wb.active  # sheet 활성화 하기
        for i in range(1, len(self.list) + 1):  # 1~리스트의 크기 + 1의 범위 까지 반복
            list1 = self.list[i - 1].split(Division)  # 리스트 0부터 division으로 입력된 값으로 글자 자르기
            for k in range(1, len(list1) + 1):  # 1~리스트1의 크기 +1 만큼 반복
                sheet[chr(k + 64) + str(i)].value = list1[k - 1]  # ASCII의 65부터 A,B,...시트 번호 A1부터 하나씩 입력한다.
        name = QFileDialog.getSaveFileName()
        wb.save(name[0] + '.xlsx')
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Load_Sheet(self, fName):  # 이용할 엑셀파일을 워크북으로 로딩하기
        self.wb = openpyxl.load_workbook(fName)  # sheet를 load
        self.ws = self.wb.active
        self.Name = fName  # 파일네임 그대로 저장

    def Delete_rows(self, row):  # 열 지우기
        rowNum = int(row)  # 입력받은 문자를 int로 바꿔주기
        self.ws.delete_rows(rowNum)  # row 지우기
        self.wb.save(self.Name)  # 파일 저장
        QMessageBox.about(self, "message", "삭제되었습니다.")

    def Delete_cols(self, col):  # 행 지우기
        colNum = int(col)  # 입력받은 문자를 int로 바꿔주기
        self.ws.delete_cols(colNum)  # col 지우기
        self.wb.save(self.Name)  # 파일저장
        QMessageBox.about(self, "message", "삭제되었습니다.")

    def Move_Cell(self, Range, move_range):  # 이동시킬 값의 범위, 얼만큼 이동할란지? ex)3,3
        word_list = move_range.split(',')  # 입력받은 범위를 ,로 split해 list에 각 문자 담기
        moverow = int(word_list[0])  # 리스트의 문자열을 숫자로 바꿔주기
        movecol = int(word_list[1])
        self.ws.move_range(Range, rows=moverow, cols=movecol, translate=True)
        self.wb.save(self.Name)
        QMessageBox.about(self, "message", "완료되었습니다.")

    def GetValue(self, cell):  # 지정범위 값 txt 파일로 적기
        div = list(cell)  # 전달받은 범위를 한글자씩 쪼개서 리스트에 담기
        cell_range = self.ws[cell]  # 입력받은 셀범위를 워크시트에서 찾기
        a = list()
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        f = open(file_name, "w", encoding="utf-8")
        if ":" in div:  # 여러셀 가져오기
            for row in cell_range:  # 셀의 개수만큼 반복
                for cell in row:  # row만큼 반복
                    a.append(cell.value)  # 리스트에 값 추가
            for i in range(len(a)):
                f.write('%s\n' % a[i])  # txt파일에 적기
        else:  # 지정한 셀만 필요시
            f.write(cell_range.value)

    def Find_row(self, col , text, totalcol):  # 지정한col에 text가 존재하는 열의 값들 모두 가져오기
        colnum = int(col)
        a = list()
        n = int(totalcol)
        for i in self.ws.rows:  # 모든 row를 검사
            if operator.eq(i[colnum - 1].value, text):  # 해당 컬럼에서 text와 같은 값이 존재하나?
                for j in range(0, n):  # 총 col개수만큼 반복
                    a.append(i[j].value)  # 리스트에 저장
        result = [a[i * n:(i + 1) * n] for i in range((len(a) + n - 1) // n)]  # 리스트에 셀의 값 하나씩 저장
        num = int(len(a) / n)   # a리스트의 크기를 총 col의 개수로 나눈 값을 num에 저장
        self.wb.close()
        wb2 = openpyxl.Workbook()  # 값 저장 할 wb 열기
        ws2 = wb2.active
        for i in range(1, num + 1):
            for j in range(0, n):
                ws2.cell(row=i, column=j + 1).value = result[i - 1][j]  # ws2의 셀에 값 입력
        name = QFileDialog.getSaveFileName()
        wb2.save(name[0] + '.xlsx')

    def SetValue(self, cell, text):  # 지정한 셀에 값 넣기
        self.ws[cell] = text
        self.wb.save(self.Name)
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Read_docx(self, fileName):  # 워드파일 내용 읽어오기
        text = docx2txt.process(fileName)  # 워드파일 읽어오기
        QMessageBox.about(self, "message", text)

    def Make_docx(self,style):
        if 'docx' in style:  # style변수에 스타일이 지정된 워드파일이 입력됐을 경우
            self.document = Document(style)  # document를 지정한 스타일로 만들기
        else:
            self.document = Document()  # document를 지정한 형식 없이 만들기

    def Add_Image(self,filename,size):
        if '_size' in size :  # size에 입력값을 안넣었을 경우
            self.document.add_picture(filename)
        else:
            size2 = size.split(',')  # 사이즈를 ,로 구분
            sizelist2 = [int(x) for x in size2]  # size를 리스트에 저장
            self.document.add_picture(filename,width=Cm(sizelist2[0]),height= Cm(sizelist2[1]))  # cm단위로 사진 붙이기

    def Write_head(self, text, style):  # 제목작성
        if '_style' in style:  # style에 아무것도 지정되지 않았을 경우
            self.document.add_heading(text, 0)  # 헤더의 크기를 기본 0 으로 지정하거나 , Make_docx에서 style을 지정했을 경우엔 위에서 지정한 값으로 지정
        else:
            numstyle = int(style)  # 입력받은 값을 int로 바꿔주기
            self.document.add_heading(text, numstyle)  # 기본 워드의 스타일 번호로 스타일을 준다.

    def Write_body(self, text, style):  # 본문작성
        if style == '_style':
            self.document.add_paragraph(text)  # style에 지정된 값이 없을 경우엔 기본으로 값 입력
        else:
            self.document.add_paragraph(text, style=style)  # style이 지정된 경우 스타일을 준다.

    def Save_docx(self):  # 워드파일 저장
        name = QFileDialog.getSaveFileName()
        self.document.save(name[0] + '.docx')

    def Read_csv(self,filename):  # 읽기전용으로 csv파일 열기
        f = open(filename, 'r', encoding='euc-kr')  # csv파일 읽기
        self.rdr = csv.reader(f)  # csv파일 리더 생성

    def Get_Line(self, selectline):
        linenum = int(selectline)
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        txt = open(file_name, "w", encoding='utf-8')  # 텍스트파일 쓰기형식으로 열기
        for line in self.rdr:
            txt.write(line[linenum] + '\n')  # 텍스트파일에 해당줄의 값 입력
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Open_Add_mode(self,filename):
        self.f = open(filename, 'a', newline='')  # 수정형식으로 파일 열기
        self.wr = csv.writer(self.f)  # writer지정

    def Add_row(self,text):
        rs2 = text.split(',')  # 입력된 값을 ,로 구분
        self.wr.writerow(rs2)  # row입력
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Write_csv(self, text):
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".csv"
        f = open(file_name, "w", encoding='euc-kr', newline='')  # 쓰기전용으로 파일열기
        wr = csv.writer(f)  # writer 지정
        rs2 = text.split(',')
        wr.writerow(rs2)

    def File_Close(self):
        self.f.close()

    def Get_All_Value(self):  # 모든 값 가져와 txt 파일로 저장
        a = list()
        name = QFileDialog.getSaveFileName()
        file_name = name[0] + ".txt"
        text = open(file_name, "w", encoding='utf-8')  # 쓰기형식으로 파일열기
        for txt in self.rdr:
            a.append(txt)  # 리스트에 텍스트 추가
        for i in range(len(a)):
            text.write('%s\n' % a[i])  # txt파일에 값입력
        QMessageBox.about(self, "message", "완료되었습니다.")

    def Close_Browser(self): # 브라우저 닫는 함수
        self.driver.close() # 현재 사용중이던 브라우저를 닫는다.

    def Close_PopUp(self): # 브라우저에 생성된 팝업 창을 닫는 함수
        driver_count = len(self.driver.window_handles) # 현재 열려 있는 창의 수를 driver_count에 저장
        if driver_count != 1: # 열려있는 창이 1개가 아니라면(팝업 창이 존재한다면)
            driver_count = driver_count - 1
            self.driver.switch_to.window(self.driver.window_handles[driver_count])  # 팝업창을 선택
            self.driver.close()
        self.driver.switch_to.window(self.driver.window_handles[0])  # main창 선택

    def Image_Download(self, src):
        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        urllib.request.urlretrieve(src, name[0] + '.png')  # src주소로 이미지를 찾아 '.png'형식으로 저장

    def Page_Screenshot(self):
        name = QFileDialog.getSaveFileName()  # 다른 이름으로 저장
        self.driver.save_screenshot(name[0] + '.png')  # 드라이버가 위치한 페이지에서 스크린샷 후 '.png'형식으로 저장

    def Image_Click(self, ImageFile):
        img = ImageFile  # 이미지 파일을 저장
        pos = imagesearch(img)  # 이미지 파일과 일치하는 이미지의 위치를 반환하는 imagersearch함수호출
        if pos[0] != -1:  # 찾았을 경우
            click_image(img, pos, 'left', 0.1)  # 이미지, 위치, 활성화할 마우스 버튼, 마우스가 원래 위치에서 새 위치로 이동하는 데 걸리는 시간 순

    def Position_Click(self, X, Y):
        pyautogui.click(int(X), int(Y))  # 함수 호출 시 좌표를 찾는 GUI.py -> selectPosition함수 호출

    def Delay(self, _time):
        time.sleep(int(_time))  # 지정한 초 동안 딜레이